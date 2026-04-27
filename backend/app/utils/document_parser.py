"""
文档解析模块
支持解析 PDF、DOC、DOCX、TXT、Markdown 格式的文档，返回纯文本内容
"""
import os
import subprocess
from pathlib import Path


class DocumentParser:
    """文档解析器"""

    @staticmethod
    def parse(file_path: str) -> str:
        """
        根据文件扩展名自动选择解析器
        :param file_path: 文件路径
        :return: 解析后的纯文本内容
        """
        suffix = Path(file_path).suffix.lower()

        parsers = {
            ".pdf": DocumentParser._parse_pdf,
            ".doc": DocumentParser._parse_doc,
            ".docx": DocumentParser._parse_docx,
            ".txt": DocumentParser._parse_txt,
            ".md": DocumentParser._parse_txt,
        }

        parser = parsers.get(suffix)
        if parser is None:
            raise ValueError(f"不支持的文件格式: {suffix}，仅支持 PDF、DOC、DOCX、TXT、Markdown")

        return parser(file_path)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """
        解析 PDF 文件
        使用 PyPDF2 提取文本
        :param file_path: PDF 文件路径
        :return: 提取的文本内容
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2")

        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts)

    @staticmethod
    def _parse_doc(file_path: str) -> str:
        """
        解析旧版 DOC 文件（OLE2 格式）
        优先使用 antiword，其次使用 LibreOffice 转换
        :param file_path: DOC 文件路径
        :return: 提取的文本内容
        """
        # 方法 1: 尝试使用 antiword
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 方法 2: 使用 LibreOffice 转换为临时文本文件
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmp_dir:
                subprocess.run(
                    [
                        "libreoffice", "--headless", "--convert-to", "txt:Text",
                        "--outdir", tmp_dir, file_path,
                    ],
                    capture_output=True,
                    timeout=60,
                )
                # 查找生成的 txt 文件
                base_name = Path(file_path).stem
                txt_path = os.path.join(tmp_dir, f"{base_name}.txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
                        return f.read().strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        raise ValueError(
            f"无法解析 .doc 文件。请将文件转换为 .docx 格式后重新上传，"
            f"或确保系统已安装 antiword 或 LibreOffice。"
        )

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """
        解析 DOCX 文件
        使用 python-docx 提取文本
        :param file_path: DOCX 文件路径
        :return: 提取的文本内容
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """
        解析纯文本文件（TXT / Markdown）
        :param file_path: 文本文件路径
        :return: 文件内容
        """
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError(f"无法解析文件 {file_path}，尝试了多种编码均失败")

    @staticmethod
    def get_supported_extensions() -> list[str]:
        """获取支持的文件扩展名列表"""
        return [".pdf", ".doc", ".docx", ".txt", ".md"]
